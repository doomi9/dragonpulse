"""Unified text extraction + chunking for the knowledge base.

Supports the document types contractors typically keep: PDF (pdfplumber),
DOCX (python-docx, optional), and plain text / Markdown. Extraction works from
raw bytes (Streamlit uploads) or a path.

Security note: we never log document text — only file names, sizes, and the
number of characters/chunks produced.
"""

from __future__ import annotations

import importlib.util
import io
import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

from dragonpulse.config.logging_config import get_logger

logger = get_logger(__name__)

_TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".rst", ".csv", ".json"}

# Sentence boundary: end punctuation followed by whitespace + a capital/digit.
_SENTENCE_RE = re.compile(r"(?<=[.!?])\s+(?=[\"'(\[]?[A-Z0-9])")
# Leading enumerators stripped from heading labels (e.g. "1.2 ", "IV. ", "# ").
_HEADING_PREFIX_RE = re.compile(r"^(#{1,6}\s*|(\d+(\.\d+)*|[IVXLCDM]+)[.)]\s+)")


class UnsupportedDocument(ValueError):
    """Raised when a document type cannot be extracted."""


class ScannedPDFError(UnsupportedDocument):
    """Raised when a PDF has no text layer (scanned/image-only) — OCR candidate."""


def format_upload_limit(max_mb: int) -> str:
    """Human-readable upload cap (e.g. 1000 -> ``1 GB``)."""
    if max_mb >= 1000 and max_mb % 1000 == 0:
        gb = max_mb // 1000
        return "1 GB" if gb == 1 else f"{gb} GB"
    return f"{max_mb} MB"


def validate_upload_size(size_bytes: int, max_mb: int, filename: str = "file") -> None:
    """Reject uploads larger than ``max_mb`` megabytes.

    Raises
    ------
    UnsupportedDocument
        If ``size_bytes`` exceeds the configured limit.
    """
    max_bytes = max_mb * 1024 * 1024
    if size_bytes > max_bytes:
        size_mb = size_bytes / (1024 * 1024)
        limit = format_upload_limit(max_mb)
        raise UnsupportedDocument(
            f"'{filename}' is {size_mb:.1f} MB — maximum upload size is {limit} per file."
        )


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """Extract plain text from uploaded ``data`` based on ``filename`` suffix.

    Raises
    ------
    UnsupportedDocument
        If the file type is not supported or extraction yields nothing.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        text, had_images = _extract_pdf_bytes(data)
    elif suffix == ".docx":
        text, had_images = _extract_docx_bytes(data), False
    elif suffix in _TEXT_SUFFIXES:
        text, had_images = data.decode("utf-8", errors="replace"), False
    else:
        raise UnsupportedDocument(
            f"Unsupported file type '{suffix}'. Supported: PDF, DOCX, TXT, MD."
        )

    text = (text or "").strip()
    if not text:
        if suffix == ".pdf" and had_images:
            raise ScannedPDFError(
                f"'{filename}' has no selectable text — it looks like a scanned / "
                "image-only PDF. Enable OCR to read it, or add a text layer first."
            )
        raise UnsupportedDocument(f"No extractable text found in '{filename}'.")
    logger.info("Extracted %d chars from %s", len(text), filename)
    return text


def extract_text_from_path(path: Path) -> str:
    """Extract text from a file on disk."""
    path = Path(path)
    return extract_text_from_bytes(path.read_bytes(), path.name)


def extract_text_with_ocr(
    data: bytes,
    filename: str,
    *,
    ocr_enabled: bool = True,
    dpi: int = 200,
    page_callback=None,
) -> str:
    """Extract text, transparently OCR'ing scanned/image-only PDFs.

    Tries normal extraction first; on a :class:`ScannedPDFError`, falls back to
    OCR when ``ocr_enabled`` and the OCR stack is available. Re-raises otherwise.
    """
    try:
        return extract_text_from_bytes(data, filename)
    except ScannedPDFError:
        if ocr_enabled and ocr_available():
            return ocr_pdf_bytes(data, dpi=dpi, page_callback=page_callback)
        raise


def ocr_available() -> bool:
    """True when the OCR stack (PyMuPDF + pytesseract + tesseract binary) is usable."""
    if importlib.util.find_spec("fitz") is None:
        return False
    if importlib.util.find_spec("pytesseract") is None:
        return False
    return shutil.which("tesseract") is not None


def ocr_pdf_bytes(
    data: bytes,
    *,
    dpi: int = 200,
    lang: str = "eng",
    page_callback=None,
) -> str:
    """OCR a scanned / image-only PDF into plain text.

    Renders each page to an image with PyMuPDF (no external binaries needed) and
    recognizes text with Tesseract via ``pytesseract``. ``page_callback(done,
    total)`` is invoked after each page so callers can show progress.

    Raises
    ------
    UnsupportedDocument
        If the OCR stack is unavailable or no text could be recognized.
    """
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image
    except ImportError as exc:  # pragma: no cover - optional deps
        raise UnsupportedDocument(
            "OCR requires PyMuPDF and pytesseract "
            "(pip install pymupdf pytesseract)."
        ) from exc

    tess = shutil.which("tesseract")
    if not tess:
        raise UnsupportedDocument(
            "OCR requires the Tesseract engine (macOS: `brew install tesseract`)."
        )
    pytesseract.pytesseract.tesseract_cmd = tess

    parts: List[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        total = doc.page_count
        for i, page in enumerate(doc, start=1):
            pix = page.get_pixmap(dpi=dpi)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            parts.append(pytesseract.image_to_string(img, lang=lang))
            if page_callback is not None:
                page_callback(i, total)

    text = "\n\n".join(parts).strip()
    if not text:
        raise UnsupportedDocument(
            "OCR produced no text — the document may be blank or unreadable."
        )
    logger.info("OCR extracted %d chars from a %d-page PDF", len(text), total)
    return text


def _extract_pdf_bytes(data: bytes) -> tuple[str, bool]:
    """Return ``(text, had_images)``.

    ``had_images`` flags scanned / image-only PDFs (no text layer) so callers can
    give an OCR-specific error instead of a generic "no text found".
    """
    import pdfplumber

    parts: List[str] = []
    had_images = False
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
            if not had_images and page.images:
                had_images = True
    return "\n\n".join(parts), had_images


def _extract_docx_bytes(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - depends on optional install
        raise UnsupportedDocument(
            "DOCX support requires 'python-docx' (pip install python-docx)."
        ) from exc

    document = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # Include table cell text too — proposals often use tables.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


@dataclass
class ChunkPiece:
    """One semantic chunk plus the nearest heading it falls under."""

    text: str
    section: Optional[str] = None


def chunk_text(
    text: str,
    *,
    chunk_chars: int = 3600,
    overlap: int = 400,
) -> List[str]:
    """Back-compat wrapper returning chunk texts (see :func:`semantic_chunks`)."""
    return [p.text for p in semantic_chunks(text, chunk_chars=chunk_chars, overlap=overlap)]


def semantic_chunks(
    text: str,
    *,
    chunk_chars: int = 3600,
    overlap: int = 400,
) -> List[ChunkPiece]:
    """Split ``text`` into coherent, overlapping chunks with section awareness.

    Improvements over naive fixed-size splitting:

    - **Heading-aware:** detected headings start a new chunk (when the current
      one already has substance) so a chunk rarely straddles two sections, and
      each chunk records the heading it belongs to.
    - **Sentence-aware:** oversized paragraphs are split on sentence boundaries
      (not mid-word), and the inter-chunk ``overlap`` carries whole trailing
      sentences for smoother context continuity.
    - **Larger targets:** defaults aim for ~900-token chunks so a strong model
      receives more complete context per retrieval.
    """
    if overlap >= chunk_chars:
        raise ValueError("overlap must be smaller than chunk_chars")

    normalized = _normalize_whitespace(text)
    if not normalized:
        return []

    blocks = _blocks_with_sections(normalized)
    pieces: List[ChunkPiece] = []
    current = ""
    start_section: Optional[str] = None

    def flush() -> None:
        nonlocal current
        if current.strip():
            pieces.append(ChunkPiece(text=current.strip(), section=start_section))
        current = ""

    for para, section in blocks:
        # Prefer to break at a heading boundary once the chunk has real substance.
        if (
            current
            and section != start_section
            and len(current) >= chunk_chars * 0.5
        ):
            tail = _overlap_tail(current, overlap)
            flush()
            current = tail
            start_section = section

        if not current:
            start_section = section

        if len(para) > chunk_chars:
            if current.strip():
                flush()
            for sub in _sentence_split(para, chunk_chars, overlap):
                pieces.append(ChunkPiece(text=sub.strip(), section=section))
            current = ""
            start_section = None
            continue

        candidate = f"{current}\n\n{para}".strip() if current else para
        if len(candidate) <= chunk_chars:
            current = candidate
        else:
            tail = _overlap_tail(current, overlap)
            flush()
            current = f"{tail}\n\n{para}".strip() if tail else para
            start_section = section

    flush()
    return pieces


def _blocks_with_sections(normalized: str) -> List["tuple[str, Optional[str]]"]:
    """Split into content paragraphs, tagging each with the active heading."""
    blocks: List[tuple] = []
    section: Optional[str] = None
    for para in normalized.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if _is_heading(para):
            section = _clean_heading(para)
            continue
        blocks.append((para, section))
    return blocks


def _is_heading(para: str) -> bool:
    """Heuristically decide whether a paragraph is a section heading."""
    if "\n" in para:
        return False
    s = para.strip()
    if not (3 <= len(s) <= 90):
        return False
    if s.startswith("#"):
        return True
    if s.endswith((".", "!", "?", ",", ";", ":")) and not s.startswith("#"):
        # A trailing colon is common on headings ("Scope:"), allow that one.
        if not s.endswith(":"):
            return False
    if re.match(r"^(\d+(\.\d+)*|[IVXLCDM]+)[.)]\s+\S", s) and not s.endswith("."):
        return True
    words = s.rstrip(":").split()
    if not (1 <= len(words) <= 12):
        return False
    letters = [c for c in s if c.isalpha()]
    if letters and sum(c.isupper() for c in letters) / len(letters) >= 0.6:
        return True  # mostly uppercase → heading
    capitalized = sum(1 for w in words if w[:1].isupper())
    return len(words) >= 2 and capitalized / len(words) >= 0.7


def _clean_heading(para: str) -> str:
    return _HEADING_PREFIX_RE.sub("", para.strip()).rstrip(":").strip()[:90]


def _split_sentences(text: str) -> List[str]:
    return [s for s in _SENTENCE_RE.split(text) if s.strip()]


def _overlap_tail(text: str, overlap: int) -> str:
    """Trailing context for the next chunk: whole sentences up to ``overlap`` chars."""
    if overlap <= 0 or not text:
        return ""
    if len(text) <= overlap:
        return text
    tail = ""
    for sentence in reversed(_split_sentences(text)):
        candidate = f"{sentence} {tail}".strip() if tail else sentence
        if len(candidate) > overlap and tail:
            break
        tail = candidate
        if len(tail) >= overlap:
            break
    if not tail:
        tail = text
    return tail[-overlap:]


def _sentence_split(text: str, chunk_chars: int, overlap: int) -> List[str]:
    """Split an oversized paragraph on sentence boundaries (char-window fallback)."""
    sentences = _split_sentences(text)
    if len(sentences) <= 1:
        return _char_windows(text, chunk_chars, overlap)
    out: List[str] = []
    current = ""
    for sentence in sentences:
        if len(sentence) > chunk_chars:
            if current:
                out.append(current)
                current = ""
            out.extend(_char_windows(sentence, chunk_chars, overlap))
            continue
        candidate = f"{current} {sentence}".strip() if current else sentence
        if len(candidate) <= chunk_chars:
            current = candidate
        else:
            out.append(current)
            tail = _overlap_tail(current, overlap)
            current = f"{tail} {sentence}".strip() if tail else sentence
    if current:
        out.append(current)
    return out


def _char_windows(text: str, chunk_chars: int, overlap: int) -> List[str]:
    step = max(1, chunk_chars - overlap)
    return [text[i : i + chunk_chars] for i in range(0, len(text), step)]


def _normalize_whitespace(text: str) -> str:
    # Collapse 3+ newlines to a paragraph break; strip trailing spaces per line.
    lines = [ln.rstrip() for ln in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    cleaned = "\n".join(lines)
    while "\n\n\n" in cleaned:
        cleaned = cleaned.replace("\n\n\n", "\n\n")
    return cleaned.strip()


# Re-export for callers that want both in one place.
def extract_and_chunk(
    data: bytes,
    filename: str,
    *,
    chunk_chars: int = 3600,
    overlap: int = 400,
) -> "tuple[str, List[str]]":
    """Extract then chunk in one call. Returns ``(full_text, chunks)``."""
    text = extract_text_from_bytes(data, filename)
    return text, chunk_text(text, chunk_chars=chunk_chars, overlap=overlap)


__all__ = [
    "UnsupportedDocument",
    "ScannedPDFError",
    "ChunkPiece",
    "extract_text_from_bytes",
    "extract_text_from_path",
    "extract_text_with_ocr",
    "ocr_available",
    "ocr_pdf_bytes",
    "chunk_text",
    "semantic_chunks",
    "extract_and_chunk",
    "format_upload_limit",
    "validate_upload_size",
]
