"""Grounded proposal draft generation.

Pipeline (per section):

1. **Solicitation context** — the opportunity's attachment text is chunked and
   embedded into an *ephemeral* in-memory index (reusing the knowledge base's
   embedding backend). A section-specific query retrieves the most relevant
   solicitation passages.
2. **Company context** — the same query retrieves the most relevant chunks from
   the persistent RAG knowledge base (past proposals/performance).
3. **Generation** — both contexts (each labeled for citation) plus opportunity
   metadata are handed to the grounded LLM, which is instructed to use only the
   provided context and cite it. If no LLM is available, a deterministic
   *evidence scaffold* is produced instead, so the feature still works offline.

Every section records the exact evidence it used, so the whole draft is
auditable and never asserts a capability that is not in the cited context.
"""

from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple

import numpy as np

from dragonpulse.config.logging_config import get_logger
from dragonpulse.config.settings import Settings, get_settings
from dragonpulse.models.knowledge import Document
from dragonpulse.models.opportunity import Opportunity
from dragonpulse.models.proposal import (
    STATUS_ADDRESSED,
    STATUS_NOT_ADDRESSED,
    STATUS_PARTIAL,
    CitationEvidence,
    ComplianceItem,
    ComplianceMatrix,
    ProposalDraft,
    ProposalSection,
)
from dragonpulse.processors.embeddings import EmbeddingBackend, get_embedding_backend
from dragonpulse.processors.knowledge_base import KnowledgeBase
from dragonpulse.processors.llm import LLMClient, LLMUnavailable
from dragonpulse.processors.text_extract import chunk_text

logger = get_logger(__name__)

_SNIPPET_CHARS = 320
# Style exemplars use a long, structure-preserving excerpt so the model can study
# the company's full section organization, headings, bullet conventions, sentence
# rhythm, and level of technical detail — not just a few keywords. These are the
# templates the new section is patterned after, so they are intentionally generous.
# templates the new section is patterned after — often a full semantic chunk.
_EXEMPLAR_CHARS = 3200
_MAX_SOLICITATION_CHARS = 60_000  # safety cap before chunking
# How many KB hits to consider before section/doc reranking for style exemplars.
_STYLE_SEARCH_POOL = 16
# Bonus added to cosine score when a chunk's section heading matches the target.
_SECTION_MATCH_BONUS = 0.18

# Proposal-specific system prompt. The knowledge base is treated as LEARNING
# MATERIAL for the company's voice/style/capabilities — not a fact database to
# insert. The model learns the company's style and then writes a NEW, relevant
# section for the current opportunity, only pulling specific KB facts when they
# directly apply (so it stops forcing in unrelated past-project details).
_PROPOSAL_SYSTEM_PROMPT = (
    "You are an expert U.S. government proposal writer drafting on behalf of the "
    "SPECIFIC company whose documents appear in the CONTEXT. Infer the company's real "
    "name and identity from its documents and write in its voice, first person plural "
    "('we', 'our team'); never call the company 'DragonPulse' (that is only the name "
    "of this drafting tool).\n\n"
    "USE THE KNOWLEDGE BASE AS LEARNING MATERIAL, NOT A FACT DATABASE: The company "
    "documents in the CONTEXT are primarily training material for the company's voice, "
    "tone, style, capabilities, and how it structures proposals. LEARN from them, then "
    "write a NEW section tailored to THIS specific opportunity. Do NOT treat the "
    "knowledge base as a set of facts that must be inserted into the proposal.\n\n"
    "REASON AND WRITE ORIGINALLY: You are a capable expert — think through what this "
    "opportunity actually needs and write a strong, original, well-reasoned section. "
    "Only a few of the most relevant knowledge-base passages are provided on purpose; "
    "treat them as cues, not a script. Synthesize and reason rather than stitching "
    "retrieved snippets together.\n\n"
    "RELEVANCE — be selective: Only reference specific details from the knowledge base "
    "if they directly apply to this opportunity. Do NOT force in unrelated technical "
    "details, projects, products, or numbers just because they appear in the knowledge "
    "base. If the knowledge base has no clearly relevant content for this section, "
    "write the section from sound industry best practices and the solicitation's "
    "requirements — still in the company's established tone and style — rather than "
    "padding it with irrelevant knowledge-base material.\n\n"
    "DO NOT FABRICATE: Facts about the requirement must come from the SOLICITATION "
    "EXCERPTS. Company-specific facts — named past contracts, customers, dollar "
    "amounts, certifications, or personnel — must come from the knowledge base; never "
    "invent them. You MAY write general methodology and approach in the company's "
    "voice without a specific citation. When you DO use specific content from the "
    "knowledge base or the solicitation, cite its bracketed source label, e.g. "
    "[KB: Capabilities #1] or [Solicitation: SOW.pdf #2].\n\n"
    "STYLE — SOUND EXACTLY LIKE THIS COMPANY: The company-writing-style excerpts are "
    "your STYLE TEMPLATE. Study them closely and emulate, as faithfully as you can: "
    "(a) STRUCTURE — how they organize a section, the order of ideas, and whether they "
    "open with a summary statement, a methodology, or a capability claim; (b) FORMAT — "
    "their heading conventions and capitalization, their bullet vs. paragraph mix, and "
    "list formatting; (c) SENTENCE STYLE — sentence length and rhythm, active vs. "
    "passive voice, and how technical they get; (d) VOCABULARY & PHRASING — the actual "
    "terms, transitions, and recurring phrases the company uses. Pattern your new "
    "section on these templates so a reader would believe the same author wrote it.\n\n"
    "AVOID GENERIC PROPOSAL LANGUAGE: Do NOT fall back on bland, boilerplate corporate "
    "phrasing (e.g. 'We are committed to excellence', 'leveraging best-in-class "
    "solutions', 'world-class team') when the company's own documents show a more "
    "specific, concrete way of saying it. Prefer adapting and rephrasing the company's "
    "real wording and level of detail over writing from scratch in a default style. "
    "Apply this learned style to NEW content that fits THIS opportunity — match the "
    "voice and structure, not the unrelated facts.\n\n"
    "CITATIONS & READABILITY: The proposal must read as polished, submission-ready "
    "prose — not an annotated bibliography. Do NOT litter every sentence with "
    "bracketed citations. Cite sparingly: at the end of a paragraph or subsection "
    "when a specific sourced fact appears, or once when introducing a block of "
    "solicitation-derived requirements. Prefer flowing professional text; full "
    "source provenance is recorded separately in the Sources panel.\n\n"
    "OUTPUT FORMAT — PLAIN TEXT ONLY: Write submission-ready plain text. Do NOT use "
    "markdown syntax (no **, ##, _, or bullet markdown). Use formal headings on their "
    "own lines (e.g. 'Tab 2 — Equipment Drawings and Descriptive Data' or "
    "'1.0 Excitation Equipment'). Do NOT append a 'Sources', 'References', or "
    "'Citations' section to your output — source tracking is handled separately.\n\n"
    "TRADITIONAL GOVERNMENT PROPOSAL FORMAT: Produce a formal Volume 1 Technical "
    "Proposal in the style of U.S. federal/USACE solicitations. Use the numbered "
    "Tab structure (Tab 1, Tab 2, …), formal headings, and professional government "
    "contracting language. Write FULL-LENGTH content — each Tab should be substantive "
    "(several printed pages). The complete draft should target 20+ pages when printed. "
    "Use numbered subsections, specification-style cross-references (§) where natural, "
    "and technical depth matching the uploaded template proposal — not a short summary "
    "or modern SaaS-style proposal."
)

# Appended to every section instruction: tie content to THIS opportunity, match
# style, and use the KB only where it is directly relevant.
_STYLE_ADDENDUM = (
    "\n\nWrite in the company's voice for THIS opportunity, using the solicitation "
    "excerpts as the basis for what to address. Use the writing-style reference and "
    "document outline as your TEMPLATE: reproduce their section flow, heading style, "
    "paragraph vs. bullet mix, sentence length and rhythm, formality level, and depth "
    "of technical detail — a reader should believe the same author wrote both documents. "
    "Adapt the company's characteristic phrasing and voice; do not invent a new generic structure. "
    "Avoid bland corporate boilerplate when the template shows a more specific voice. "
    "Use knowledge-base specifics ONLY where they directly apply; otherwise write from "
    "sound practices in the company's style. Cite sources sparingly so the prose reads "
    "naturally — do not break the flow with a citation on every sentence.\n\n"
    "LENGTH: Write a FULL, submission-ready section for this Tab — multiple pages of "
    "content with detailed subsections. Do not abbreviate or summarize; match the depth "
    "and page count of the corresponding Tab in the template proposal."
)

# Retrieval queries that surface the compliance-relevant parts of a solicitation.
_COMPLIANCE_QUERIES = [
    "Section L instructions to offerors proposal submission requirements format page limit",
    "Section M evaluation criteria factors basis for award technical evaluation",
    "the contractor shall perform mandatory tasks requirements deliverables",
    "offeror must submit provide required documentation certifications registration",
]
_VALID_CATEGORIES = ("Section L", "Section M", "Evaluation", "SOW (shall)", "Other")
# Sentence-level detector for mandatory/evaluation language (rules fallback).
_REQUIREMENT_RE = re.compile(
    r"[^.\n]*\b(shall|must|will be evaluated|is required to|are required to|"
    r"offeror[s]? (?:shall|must|will)|the government will evaluate)\b[^.\n]*\.",
    re.IGNORECASE,
)
# Status thresholds on cosine similarity between a requirement and its best
# matching proposal section. Tuned for L2-normalized semantic embeddings; these
# are intentionally conservative so the auto-status under-claims rather than
# over-claims coverage. They are only a starting point the user edits in the UI.
_ADDRESSED_THRESHOLD = 0.58
_PARTIAL_THRESHOLD = 0.45
# Below this best-KB-cosine, treat direct past-performance matches as "weak" and
# steer the model toward an honest transferable-experience framing.
_KB_WEAK_THRESHOLD = 0.45
# KB fact chunks scoring below this cosine are dropped from the "optional
# reference" context entirely — they are not relevant enough to risk the model
# forcing them in. (Tuned for semantic embeddings; with the lexical hashing
# backend scores are compressed, so the prompt's relevance rules do the heavy
# lifting and this acts as a light floor.)
_KB_RELEVANCE_FLOOR = 0.30


@dataclass
class SectionSpec:
    """Definition of a proposal section (Tab) to generate."""

    section_id: str
    title: str          # display title, e.g. "Tab 2 Equipment Drawings and Descriptive Data"
    query: str          # retrieval query (semantic) for solicitation + KB
    instruction: str    # what to write
    optional: bool = False
    deterministic: bool = False  # title page / TOC — built from template + metadata
    max_tokens: int = 2500     # per-section LLM budget (larger Tabs get more)
    sol_k: int = 5
    kb_k: int = 3
    honest_transferable: bool = False
    style_categories: Tuple[str, ...] = ()
    style_k: int = 4
    style_section_keywords: Tuple[str, ...] = ()


# Default Tab layout mirrors the Libby Dam Exciter Replacement technical volume.
_DEFAULT_TABS: List[Tuple[str, str]] = [
    ("1", "Executive Summary"),
    ("2", "Equipment Drawings and Descriptive Data"),
    ("3", "Installation Plan"),
    ("4", "Technical Comments and Clarifications"),
    ("5", "Past Performance"),
    ("6", "Appendices"),
]

_DEFAULT_APPENDICES: List[Tuple[str, str]] = [
    ("1", "Equipment Highlights / Product Description"),
    ("2", "Specification Guide / Technical Data"),
    ("3", "One-Line Diagram"),
    ("4", "Equipment Outline Drawing"),
    ("5", "Technical Standards and Certifications"),
    ("6", "HMI / Control System Information"),
    ("7", "Selected Screen Shots / Interface Examples"),
    ("8", "Quality Certification & Facility Capability"),
    ("9", "Product Description and Maintenance Procedures"),
    ("10", "Maintenance Schedule"),
    ("11", "Spare Parts List"),
]

_TAB_RE = re.compile(r"(?im)^Tab\s+(\d+)\s+(.+?)\s*$")
_APPENDIX_RE = re.compile(r"(?im)^Appendix\s+(\d+)\s+(.+?)\s*$")

_RESTRICTION_NOTICE = (
    "This proposal includes data that shall not be disclosed outside the Government "
    "and shall not be duplicated, used, or disclosed — in whole or in part — for any "
    "purpose other than to evaluate this proposal. If, however, a contract is awarded "
    "to the offeror as a result of — or in connection with — the submission of this "
    "data, the Government shall have the right to duplicate, use, or disclose the data "
    "to the extent provided in the resulting contract. This restriction does not limit "
    "the Government's right to use information contained in this data if it is obtained "
    "from another source without restriction."
)

# Order matters: this is the draft's section order (traditional Tab structure).
SECTION_SPECS: List[SectionSpec] = [
    SectionSpec(
        section_id="title_page",
        title="Title Page",
        query="",
        instruction="",
        deterministic=True,
    ),
    SectionSpec(
        section_id="table_of_contents",
        title="Table of Contents",
        query="",
        instruction="",
        deterministic=True,
    ),
    SectionSpec(
        section_id="tab_1_executive_summary",
        title="Tab 1 — Executive Summary",
        query="executive summary objective scope purpose requirement mission understanding",
        instruction=(
            "Write **Tab 1 — Executive Summary** as the opening of a formal Volume 1 "
            "Technical Proposal. Open with a courteous introduction to the agency, state "
            "understanding of the requirement, summarize the technical approach and key "
            "equipment, note subcontractor/partner arrangements if grounded in the "
            "knowledge base, and close with a confident capability statement. Write "
            "4–6 substantive paragraphs in formal government contracting language."
        ),
        max_tokens=2800,
        style_categories=("Technical", "Past Performance", "Capabilities"),
        style_section_keywords=("executive summary", "summary", "introduction"),
        style_k=4,
    ),
    SectionSpec(
        section_id="tab_2_equipment_descriptive",
        title="Tab 2 — Equipment Drawings and Descriptive Data",
        query="equipment drawings descriptive data excitation system generator voltage "
        "transformer switchgear protection controls cubicle specification",
        instruction=(
            "Write **Tab 2 — Equipment Drawings and Descriptive Data** — the largest "
            "technical Tab. Organize with formal numbered subsections matching the "
            "template (e.g. Excitation Equipment, Medium Voltage Primary AC Circuit, "
            "Low Voltage DC Circuit, Generator Characteristic Data, Protective Functions, "
            "Factory Tests). For each subsystem provide detailed descriptive paragraphs: "
            "equipment proposed, how it meets solicitation requirements, technical "
            "specifications, installation interfaces, and testing. Reference drawing "
            "numbers/appendices where appropriate (e.g. 'see Appendix 4'). This Tab "
            "should be highly detailed — the core of the technical volume."
        ),
        max_tokens=5500,
        kb_k=4,
        style_categories=("Technical", "Past Performance"),
        style_section_keywords=(
            "equipment", "excitation", "generator", "voltage", "transformer",
            "switchgear", "protection", "system description", "factory test",
        ),
        style_k=5,
        sol_k=6,
    ),
    SectionSpec(
        section_id="tab_3_installation_plan",
        title="Tab 3 — Installation Plan",
        query="installation plan removal legacy equipment wiring commissioning schedule "
        "site work outage coordination safety",
        instruction=(
            "Write **Tab 3 — Installation Plan** covering removal of existing equipment, "
            "installation and wiring of new systems, site coordination, outage planning, "
            "safety procedures, testing sequence, and commissioning. Use formal subsections "
            "with step-by-step narrative in the company's established installation style. "
            "Include subcontractor roles if grounded in the knowledge base."
        ),
        max_tokens=4200,
        style_categories=("Technical", "Past Performance"),
        style_section_keywords=("installation", "commissioning", "removal", "wiring"),
        style_k=4,
        sol_k=5,
    ),
    SectionSpec(
        section_id="tab_4_technical_comments",
        title="Tab 4 — Technical Comments and Clarifications",
        query="technical comments clarifications exceptions deviations specification "
        "hazardous materials notices warranty testing requirements",
        instruction=(
            "Write **Tab 4 — Technical Comments and Clarifications** addressing "
            "specification interpretation, technical clarifications, exceptions or "
            "deviations (if any), notices regarding hazardous materials, warranty "
            "terms, and responses to ambiguous requirements. Mirror the formal notice "
            "style of the template (e.g. 'Notice Regarding…'). Be explicit and technical."
        ),
        max_tokens=4200,
        kb_k=4,
        style_categories=("Technical", "Past Performance"),
        style_section_keywords=(
            "notice", "clarification", "comment", "hazardous", "warranty", "exception",
        ),
        style_k=4,
        sol_k=5,
    ),
    SectionSpec(
        section_id="tab_5_past_performance",
        title="Tab 5 — Past Performance",
        query="past performance relevant experience similar projects prior contracts "
        "customers scope outcomes exciter replacement hydroelectric",
        instruction=(
            "Write **Tab 5 — Past Performance** presenting the company's relevant "
            "contract experience. For each reference provide: customer, project scope, "
            "relevance to this solicitation, outcomes, and contact/reference availability. "
            "Use the formal reference format from the template. Do NOT invent projects."
        ),
        max_tokens=3200,
        honest_transferable=True,
        style_categories=("Past Performance", "Technical"),
        style_section_keywords=("past performance", "reference", "experience", "relevant"),
        style_k=4,
    ),
    SectionSpec(
        section_id="tab_6_appendices",
        title="Tab 6 — Appendices",
        query="appendix drawings diagram specification guide product description "
        "maintenance spare parts certification highlights",
        instruction=(
            "Write **Tab 6 — Appendices** listing and describing each appendix that "
            "accompanies this technical proposal. For each appendix entry, provide a "
            "short formal description of what the appendix contains and how it supports "
            "the proposal (e.g. product highlights document, one-line diagram, equipment "
            "outline drawing, maintenance procedures, spare parts list). Note that actual "
            "drawings/documents would be attached separately; describe what would be "
            "included based on the template and solicitation."
        ),
        max_tokens=3800,
        style_categories=("Technical", "Past Performance"),
        style_section_keywords=("appendix", "appendices", "drawing", "diagram"),
        style_k=3,
    ),
]


def _snippet(text: str, limit: int = _SNIPPET_CHARS) -> str:
    text = " ".join(text.split())
    return text[:limit] + ("…" if len(text) > limit else "")


def _exemplar_snippet(text: str) -> str:
    """A long excerpt that preserves the document's structure/formatting cues.

    Newlines are kept so headings and bullets survive (they are exactly the
    structural patterns we want the model to mirror). When the excerpt must be
    trimmed, it cuts at the nearest paragraph or sentence boundary so the model
    sees whole, well-formed structures rather than a mid-sentence fragment.
    """
    cleaned = "\n".join(line.rstrip() for line in text.strip().splitlines())
    if len(cleaned) <= _EXEMPLAR_CHARS:
        return cleaned
    window = cleaned[:_EXEMPLAR_CHARS]
    # Prefer a paragraph break, then a sentence end, then a word boundary.
    cut = window.rfind("\n\n")
    if cut < int(_EXEMPLAR_CHARS * 0.6):
        sentence = max(window.rfind(". "), window.rfind(".\n"), window.rfind("\n"))
        cut = sentence if sentence >= int(_EXEMPLAR_CHARS * 0.6) else -1
    if cut == -1:
        cut = window.rfind(" ")
    if cut <= 0:
        cut = _EXEMPLAR_CHARS
    return cleaned[:cut].rstrip() + "…"


def _primary_proposal_document(kb: KnowledgeBase) -> Optional[Document]:
    """Pick the uploaded document that best represents a full proposal template.

    Prefers large Technical / Past Performance volumes (e.g. a past technical
    proposal) over line cards, capability one-pagers, or pricing forms.
    """
    candidates: List[tuple] = []
    for doc in kb.list_documents():
        score = 0.0
        name_low = doc.name.lower()
        if doc.doc_type == "Proposal":
            score += 12.0
        if doc.doc_type == "Technical Document":
            score += 8.0
        if doc.doc_type == "Past Performance" and doc.chunk_count >= 5:
            score += 10.0
        if doc.category in ("Technical", "Past Performance"):
            score += 4.0
        score += min(doc.chunk_count, 25) * 0.5
        if any(w in name_low for w in ("proposal", "excitation", "technical volume", "volume 1")):
            score += 6.0
        if doc.doc_type in ("Capability Statement", "Pricing") and doc.chunk_count <= 3:
            score -= 8.0
        if "line card" in name_low or "capabilities statement" in name_low:
            score -= 6.0
        if score > 0:
            candidates.append((score, doc))
    if not candidates:
        return None
    candidates.sort(key=lambda item: item[0], reverse=True)
    return candidates[0][1]


def _proposal_structure_outline(kb: KnowledgeBase, doc: Document) -> str:
    """Extract the section-heading flow from an uploaded proposal for mirroring."""
    chunks = sorted(
        (c for c in kb._chunks if c.doc_id == doc.doc_id),
        key=lambda c: c.ordinal,
    )
    headings: List[str] = []
    seen: set = set()
    for chunk in chunks:
        heading = (chunk.section or "").strip()
        if not heading or heading in seen:
            continue
        seen.add(heading)
        headings.append(heading)
    if not headings:
        return ""
    lines = [f"- {h}" for h in headings[:30]]
    return (
        f"Uploaded proposal template: “{doc.name}” — mirror this document's overall "
        f"section flow, heading conventions, and level of detail:\n" + "\n".join(lines)
    )


def _extract_template_tabs(text: str) -> List[Tuple[str, str]]:
    """Parse ``Tab N Title`` lines from an uploaded proposal template."""
    tabs: List[Tuple[str, str]] = []
    seen: set = set()
    for num, title in _TAB_RE.findall(text or ""):
        key = (num, title.strip())
        if key not in seen:
            seen.add(key)
            tabs.append(key)
    return tabs


def _extract_template_appendices(text: str) -> List[Tuple[str, str]]:
    """Parse ``Appendix N Title`` lines from an uploaded proposal template."""
    apps: List[Tuple[str, str]] = []
    seen: set = set()
    for num, title in _APPENDIX_RE.findall(text or ""):
        key = (num, title.strip())
        if key not in seen:
            seen.add(key)
            apps.append(key)
    return apps


def _infer_company_name(kb: KnowledgeBase, doc: Optional[Document]) -> str:
    """Best-effort company name from the template proposal's title page."""
    if doc:
        for line in kb._document_text(doc).splitlines()[:12]:
            line = line.strip()
            if not line or len(line) < 4:
                continue
            low = line.lower()
            skip_prefixes = (
                "title page", "volume", "technical proposal", "to:", "subj:", "ref:",
            )
            if low.startswith(skip_prefixes):
                continue
            if re.search(r"\b(LLC|L\.L\.C\.|INC|CORP|LTD)\b", line, re.I):
                return line
    return "The Offeror"


def _section_match_bonus(section: Optional[str], keywords: Tuple[str, ...]) -> float:
    """Score boost when a chunk's stored heading matches this section's keywords."""
    sec = (section or "").strip().lower()
    if not sec or not keywords:
        return 0.0
    for kw in keywords:
        if kw.lower() in sec:
            return _SECTION_MATCH_BONUS
    return 0.0


def _rank_style_hits(
    hits: List,
    *,
    template_doc_id: Optional[str],
    section_keywords: Tuple[str, ...],
    limit: int,
) -> List:
    """Rerank KB hits to favor the template proposal and matching section headings."""

    def sort_key(hit):
        bonus = _section_match_bonus(getattr(hit.chunk, "section", None), section_keywords)
        if template_doc_id and getattr(hit.chunk, "doc_id", None) == template_doc_id:
            bonus += 0.12
        return hit.score + bonus

    ranked = sorted(hits, key=sort_key, reverse=True)
    # De-duplicate by chunk id while preserving order.
    seen: set = set()
    out = []
    for hit in ranked:
        cid = getattr(hit.chunk, "chunk_id", None) or id(hit.chunk)
        if cid in seen:
            continue
        seen.add(cid)
        out.append(hit)
        if len(out) >= limit:
            break
    return out


def _parse_json_object(text: str) -> Dict:
    """Best-effort parse of an LLM JSON reply (handles stray prose/fences)."""
    if not text:
        return {}
    try:
        return json.loads(text)
    except (json.JSONDecodeError, TypeError):
        pass
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(text[start : end + 1])
        except json.JSONDecodeError:
            return {}
    return {}


def _normalize_category(value: Optional[str]) -> str:
    if not value:
        return "Other"
    v = str(value).strip()
    for cat in _VALID_CATEGORIES:
        if v.lower() == cat.lower():
            return cat
    return _categorize(v)


def _categorize(text: str) -> str:
    """Heuristic category from requirement text."""
    low = text.lower()
    if "section l" in low or "instructions to offeror" in low or "submission" in low:
        return "Section L"
    if "section m" in low or "basis for award" in low:
        return "Section M"
    if "evaluat" in low:
        return "Evaluation"
    if "shall" in low or "must" in low:
        return "SOW (shall)"
    return "Other"


def _prioritize(items: List["ComplianceItem"]) -> List["ComplianceItem"]:
    """Order rules-extracted items so evaluation/Section M surface first."""
    order = {"Section M": 0, "Evaluation": 1, "Section L": 2, "SOW (shall)": 3, "Other": 4}
    return sorted(items, key=lambda it: order.get(it.category, 9))


def kb_evidence_is_weak(
    evidence: List[CitationEvidence], threshold: float = _KB_WEAK_THRESHOLD
) -> bool:
    """True when the strongest knowledge-base match is below ``threshold``.

    Used to decide whether the Past Performance section should pivot to an honest
    "transferable experience" framing instead of claiming direct matches.
    """
    # Only true company *fact* chunks count here — style exemplars are learning
    # material, not evidence of relevant past performance.
    kb_scores = [
        e.score for e in evidence if e.origin == "knowledge_base" and e.score is not None
    ]
    if not kb_scores:
        return True
    return max(kb_scores) < threshold


class SolicitationIndex:
    """Ephemeral, in-memory semantic index over the solicitation attachments."""

    def __init__(self, backend: EmbeddingBackend) -> None:
        self.backend = backend
        self.chunks: List[str] = []
        self.labels: List[str] = []
        self.by_label: Dict[str, str] = {}
        self.vectors: np.ndarray = np.zeros((0, backend.dimension), dtype=np.float32)

    @property
    def is_empty(self) -> bool:
        return self.vectors.shape[0] == 0

    def add(self, attachments: List[Tuple[str, str]], chunk_chars: int, overlap: int) -> int:
        """Add ``(filename, text)`` pairs. Returns total chunks indexed."""
        all_chunks: List[str] = []
        all_labels: List[str] = []
        for name, text in attachments:
            if not text:
                continue
            text = text[:_MAX_SOLICITATION_CHARS]
            pieces = chunk_text(text, chunk_chars=chunk_chars, overlap=overlap)
            for i, piece in enumerate(pieces):
                all_chunks.append(piece)
                all_labels.append(f"Solicitation: {name} #{i + 1}")
        if not all_chunks:
            return 0
        self.chunks = all_chunks
        self.labels = all_labels
        self.by_label = dict(zip(all_labels, all_chunks))
        self.vectors = self.backend.embed(all_chunks)
        logger.info("Solicitation index built: %d chunks", len(all_chunks))
        return len(all_chunks)

    def search(self, query: str, k: int = 4) -> List[CitationEvidence]:
        if self.is_empty or not query.strip():
            return []
        q = self.backend.embed([query])[0]
        # Vectors are L2-normalized, so this dot product is cosine similarity.
        # errstate guards against spurious FP warnings from macOS Accelerate BLAS.
        with np.errstate(divide="ignore", over="ignore", invalid="ignore"):
            scores = self.vectors @ q
        scores = np.nan_to_num(scores, nan=-1.0, posinf=-1.0, neginf=-1.0)
        top = np.argsort(scores)[::-1][:k]
        return [
            CitationEvidence(
                label=self.labels[i],
                origin="solicitation",
                snippet=_snippet(self.chunks[i]),
                score=float(scores[i]),
            )
            for i in top
        ]


class ProposalGenerator:
    """Generates grounded, cited proposal sections for an opportunity."""

    def __init__(
        self,
        opportunity: Opportunity,
        *,
        settings: Optional[Settings] = None,
        knowledge_base: Optional[KnowledgeBase] = None,
        llm: Optional[LLMClient] = None,
        backend: Optional[EmbeddingBackend] = None,
    ) -> None:
        self.opp = opportunity
        self.settings = settings or get_settings()
        self.kb = knowledge_base or KnowledgeBase(self.settings)
        self.llm = llm or LLMClient(self.settings)
        self.backend = backend or self.kb.backend or get_embedding_backend(self.settings)
        self.solicitation = SolicitationIndex(self.backend)
        self._template_doc: Optional[Document] = None
        self._structure_outline: Optional[str] = None
        self._template_tabs: Optional[List[Tuple[str, str]]] = None
        self._template_appendices: Optional[List[Tuple[str, str]]] = None
        self._cached_company_name: Optional[str] = None

    def _proposal_template(self) -> Optional[Document]:
        if self._template_doc is None:
            self._template_doc = _primary_proposal_document(self.kb)
        return self._template_doc

    def _structure_outline_block(self) -> str:
        if self._structure_outline is None:
            doc = self._proposal_template()
            self._structure_outline = _proposal_structure_outline(self.kb, doc) if doc else ""
        return self._structure_outline or ""

    def _company_name(self) -> str:
        if self._cached_company_name is None:
            self._cached_company_name = _infer_company_name(
                self.kb, self._proposal_template()
            )
        return self._cached_company_name

    def _tab_structure(self) -> List[Tuple[str, str]]:
        if self._template_tabs is None:
            doc = self._proposal_template()
            text = self.kb._document_text(doc) if doc else ""
            tabs = _extract_template_tabs(text)
            self._template_tabs = tabs if tabs else list(_DEFAULT_TABS)
        return self._template_tabs

    def _appendix_structure(self) -> List[Tuple[str, str]]:
        if self._template_appendices is None:
            doc = self._proposal_template()
            text = self.kb._document_text(doc) if doc else ""
            apps = _extract_template_appendices(text)
            self._template_appendices = apps if apps else list(_DEFAULT_APPENDICES)
        return self._template_appendices

    def _render_title_page(self) -> str:
        opp = self.opp
        agency = opp.agency or opp.office or "Contracting Officer"
        subject = opp.title or "Technical Proposal"
        ref = opp.solicitation_number or opp.notice_id
        return "\n".join(
            [
                self._company_name(),
                "",
                "Title Page",
                "",
                "Volume 1",
                "",
                "Technical Proposal",
                "",
                f"To: {agency}",
                f"Subj: {subject}",
                f"Ref: {ref}",
                "",
                _RESTRICTION_NOTICE,
            ]
        )

    def _render_table_of_contents(self) -> str:
        lines = ["Table of Contents", ""]
        for num, title in self._tab_structure():
            lines.append(f"Tab {num} {title}")
        lines.append("")
        for num, title in self._appendix_structure():
            lines.append(f"Appendix {num} {title}")
        return "\n".join(lines)

    # ------------------------------------------------------------------ #
    # Context
    # ------------------------------------------------------------------ #
    def load_solicitation(self, attachments: List[Tuple[str, str]]) -> int:
        """Build the ephemeral solicitation index from extracted attachments."""
        return self.solicitation.add(
            attachments,
            chunk_chars=self.settings.rag_chunk_chars,
            overlap=self.settings.rag_chunk_overlap,
        )

    def _opportunity_block(self) -> str:
        opp = self.opp
        deadline = opp.response_deadline
        return "\n".join(
            [
                "== OPPORTUNITY METADATA ==",
                f"Title: {opp.title or 'N/A'}",
                f"Notice ID: {opp.notice_id}",
                f"Solicitation #: {opp.solicitation_number or 'N/A'}",
                f"Agency: {opp.agency or 'N/A'} / Office: {opp.office or 'N/A'}",
                f"NAICS: {opp.naics_code or 'N/A'}",
                f"Set-aside: {opp.set_aside_description or 'None'}",
                f"Response deadline: {deadline.strftime('%Y-%m-%d') if deadline else 'N/A'}",
            ]
        )

    def _gather_evidence(self, spec: SectionSpec) -> List[CitationEvidence]:
        sol = self.solicitation.search(spec.query, k=spec.sol_k)
        # Gather style exemplars first; a chunk shown as a (longer) style exemplar
        # is not repeated as a plain KB fact, so each chunk appears once.
        exemplars = self._gather_style_exemplars(spec)
        exemplar_citations = {e.label[len("KB style: "):] for e in exemplars}
        kb_hits = self.kb.search(spec.query, k=spec.kb_k)
        # Only keep fact chunks that clear the relevance floor and aren't already
        # serving as style exemplars. Low-relevance chunks are dropped so the model
        # isn't tempted to force unrelated content into the section.
        kb_ev = [
            CitationEvidence(
                label=f"KB: {h.citation}",
                origin="knowledge_base",
                snippet=_snippet(h.chunk.text),
                score=h.score,
            )
            for h in kb_hits
            if h.citation not in exemplar_citations
            and (h.score is None or h.score >= _KB_RELEVANCE_FLOOR)
        ]
        return sol + kb_ev + exemplars

    def _gather_style_exemplars(self, spec: SectionSpec) -> List[CitationEvidence]:
        """Pull writing-style exemplars from the best uploaded proposal template.

        Prefers chunks from the primary proposal document whose section headings
        match this draft section (e.g. EXECUTIVE SUMMARY, PAST PERFORMANCE), so
        the model sees how the company actually structured and wrote that part.
        """
        if spec.style_k <= 0:
            return []
        template = self._proposal_template()
        cats = list(spec.style_categories) or None
        pool = max(_STYLE_SEARCH_POOL, spec.style_k * 3)
        hits = self.kb.search(spec.query, k=pool, categories=cats)
        if not hits and cats is not None:
            hits = self.kb.search(spec.query, k=pool)
        if template:
            # Also pull from the template doc directly so section-keyword matches
            # are not missed when category filtering is narrow.
            template_hits = [
                h
                for h in self.kb.search(spec.query, k=pool)
                if getattr(h.chunk, "doc_id", None) == template.doc_id
            ]
            seen = {getattr(h.chunk, "chunk_id", id(h.chunk)) for h in hits}
            hits = list(hits) + [
                h
                for h in template_hits
                if getattr(h.chunk, "chunk_id", id(h.chunk)) not in seen
            ]
        hits = _rank_style_hits(
            hits,
            template_doc_id=template.doc_id if template else None,
            section_keywords=spec.style_section_keywords,
            limit=spec.style_k,
        )
        return [
            CitationEvidence(
                label=f"KB style: {getattr(h, 'citation', None) or h.chunk.citation()}",
                origin="style_exemplar",
                snippet=_exemplar_snippet(getattr(h.chunk, "text", "") or ""),
                score=h.score,
            )
            for h in hits
        ]

    def _context_block(self, evidence: List[CitationEvidence]) -> str:
        sol = [e for e in evidence if e.origin == "solicitation"]
        kb = [e for e in evidence if e.origin == "knowledge_base"]
        style = [e for e in evidence if e.origin == "style_exemplar"]
        parts: List[str] = []
        if sol:
            parts.append("== SOLICITATION EXCERPTS (what THIS proposal must address) ==")
            parts += [f"[{e.label}]\n{e.snippet}" for e in sol]
        if style:
            outline = self._structure_outline_block()
            parts.append(
                "\n== COMPANY WRITING STYLE — PROPOSAL TEMPLATES TO EMULATE (mirror "
                "structure & tone, not unrelated facts) =="
            )
            if outline:
                parts.append(outline)
            parts.append(
                "(The excerpts below are from the company's OWN past proposals. Treat "
                "them as the template for HOW to write this section: reproduce their "
                "section structure and ordering, heading conventions, paragraph vs. "
                "bullet style, sentence rhythm, formality, technical depth, and "
                "characteristic phrasing. The finished section should read as though it "
                "belongs in the same document set. Borrow the FORM; write NEW content "
                "for this notice. Cite sources sparingly so the prose flows naturally.)"
            )
            parts += [f"[{e.label}]\n{e.snippet}" for e in style]
        if kb:
            parts.append(
                "\n== COMPANY KNOWLEDGE BASE — OPTIONAL REFERENCE (use ONLY if directly "
                "relevant) =="
                "\n(These passages may or may not relate to this opportunity. Use a "
                "detail only if it directly applies to this solicitation, cited by "
                "label; otherwise ignore it. Do NOT force unrelated content into the "
                "section.)"
            )
            parts += [f"[{e.label}]\n{e.snippet}" for e in kb]
        if not parts:
            parts.append(
                "(No grounding context retrieved. Write this section from the "
                "solicitation requirements and industry best practices, in the "
                "company's voice.)"
            )
        return "\n\n".join(parts)

    # ------------------------------------------------------------------ #
    # Generation
    # ------------------------------------------------------------------ #
    def _tab_context_hint(self, spec: SectionSpec) -> str:
        """Remind the model which Tab it is writing within the full volume."""
        if not spec.section_id.startswith("tab_"):
            return ""
        tabs = self._tab_structure()
        lines = [
            "== VOLUME TAB STRUCTURE (mirror this organization) ==",
            *(f"Tab {num} {title}" for num, title in tabs),
        ]
        apps = self._appendix_structure()
        if apps:
            lines.append("")
            lines.append("Appendices referenced from Tab 6:")
            lines += [f"Appendix {num} {title}" for num, title in apps]
        lines.append("")
        lines.append(f"You are writing: {spec.title}")
        return "\n".join(lines)

    def generate_section(
        self,
        spec: SectionSpec,
        *,
        feedback: Optional[str] = None,
        prior: Optional[ProposalSection] = None,
    ) -> ProposalSection:
        """Generate (or regenerate) a single section, grounded and cited."""
        if spec.deterministic:
            if spec.section_id == "title_page":
                content = self._render_title_page()
            elif spec.section_id == "table_of_contents":
                content = self._render_table_of_contents()
            else:
                content = ""
            evidence: List[CitationEvidence] = []
            doc = self._proposal_template()
            if doc:
                evidence.append(
                    CitationEvidence(
                        label=f"KB style: {doc.name}",
                        origin="style_exemplar",
                        snippet="Proposal template (title page / table of contents structure)",
                    )
                )
            return ProposalSection(
                section_id=spec.section_id,
                title=spec.title,
                content=content,
                sources=evidence,
                used_llm=False,
                optional=spec.optional,
                feedback_history=list(prior.feedback_history) if prior else [],
            )

        evidence = self._gather_evidence(spec)
        tab_hint = self._tab_context_hint(spec)
        context = f"{self._opportunity_block()}\n\n"
        if tab_hint:
            context += f"{tab_hint}\n\n"
        context += self._context_block(evidence)

        instruction = spec.instruction + _STYLE_ADDENDUM
        if spec.honest_transferable and kb_evidence_is_weak(evidence):
            instruction += (
                "\n\nIMPORTANT — the knowledge base has LIMITED directly-matching "
                "past performance for this requirement. Be honest about that: do not "
                "overstate or fabricate direct experience. Instead, identify the most "
                "RELEVANT and TRANSFERABLE capabilities/projects from the provided "
                "context and explain specifically how that experience transfers to "
                "this requirement. If coverage is genuinely thin, say so plainly and "
                "frame it as transferable/adjacent experience rather than a direct match."
            )
        if feedback:
            instruction += (
                f"\n\nUSER FEEDBACK for this revision: {feedback}\n"
                "Apply the feedback while staying grounded in the CONTEXT."
            )
        if prior and prior.content:
            instruction += (
                "\n\nThe previous draft of this section is below; improve it per the "
                f"feedback rather than starting over:\n---\n{prior.content[:2000]}\n---"
            )

        used_llm = False
        content: str
        if self.llm.available:
            try:
                result = self.llm.complete(
                    instruction=instruction,
                    context=context,
                    sources=[e.label for e in evidence],
                    max_tokens=spec.max_tokens,
                    system_prompt=_PROPOSAL_SYSTEM_PROMPT,
                )
                content = result.text.strip() or self._scaffold(spec, evidence)
                used_llm = bool(result.text.strip())
            except LLMUnavailable as exc:
                logger.info("LLM unavailable for %s: %s", spec.section_id, exc)
                content = self._scaffold(spec, evidence)
        else:
            content = self._scaffold(spec, evidence)

        history = list(prior.feedback_history) if prior else []
        if feedback:
            history.append(feedback)

        content = clean_proposal_content(content)

        return ProposalSection(
            section_id=spec.section_id,
            title=spec.title,
            content=content,
            sources=evidence,
            used_llm=used_llm,
            optional=spec.optional,
            feedback_history=history,
        )

    @staticmethod
    def _scaffold(spec: SectionSpec, evidence: List[CitationEvidence]) -> str:
        """Deterministic, grounded scaffold used when no LLM is available."""
        lines = [
            f"Draft scaffold for {spec.title} — enable a local LLM "
            "(e.g. Ollama) for full prose. The grounded evidence below is what the "
            "section would be written from; nothing here is invented.",
            "",
            f"What to address: {spec.instruction}",
            "",
        ]
        sol = [e for e in evidence if e.origin == "solicitation"]
        kb = [e for e in evidence if e.origin == "knowledge_base"]
        style = [e for e in evidence if e.origin == "style_exemplar"]
        if spec.honest_transferable and kb_evidence_is_weak(evidence):
            lines.append(
                "Note: direct past-performance matches are limited. Frame the most "
                "relevant items below as transferable experience and be honest "
                "about coverage rather than claiming an exact match."
            )
            lines.append("")
        if sol:
            lines.append("Relevant solicitation requirements:")
            lines += [f"- {e.label}: {e.snippet}" for e in sol]
            lines.append("")
        if style:
            lines.append(
                "Writing-style references (learn the company's voice/structure — "
                "do not copy unrelated facts):"
            )
            lines += [f"- {e.label}: {_snippet(e.snippet)}" for e in style]
            lines.append("")
        if kb:
            lines.append(
                "Potentially relevant company evidence (use ONLY if it directly "
                "applies to this opportunity):"
            )
            lines += [f"- {e.label}: {e.snippet}" for e in kb]
            lines.append("")
        if not sol and not kb and not style:
            lines.append(
                "No grounding context retrieved. Load solicitation attachments "
                "and index relevant company documents, then regenerate."
            )
        return clean_proposal_content("\n".join(lines))

    def generate_draft(self, *, include_optional: bool = False) -> ProposalDraft:
        """Generate all sections into a :class:`ProposalDraft`."""
        specs = [s for s in SECTION_SPECS if include_optional or not s.optional]
        sections = [self.generate_section(spec) for spec in specs]
        return ProposalDraft(
            notice_id=self.opp.notice_id,
            opportunity_title=self.opp.title or self.opp.notice_id,
            agency=self.opp.agency,
            solicitation_number=self.opp.solicitation_number,
            sections=sections,
            llm_model=self.settings.llm_model if self.llm.available else None,
        )

    def strengthen_section_for_requirement(
        self,
        section: ProposalSection,
        requirement: str,
        *,
        source_label: Optional[str] = None,
    ) -> ProposalSection:
        """Regenerate ``section`` so it better addresses a specific requirement.

        Keeps the section's existing substance (passed as ``prior``) and adds
        targeted feedback, so the rest of the draft is untouched.
        """
        spec = next((s for s in SECTION_SPECS if s.section_id == section.section_id), None)
        if spec is None:
            return section
        src = f" (see {source_label})" if source_label else ""
        feedback = (
            "Strengthen this section so it explicitly and unmistakably addresses "
            f'the following solicitation requirement{src}: "{requirement.strip()}". '
            "A reviewer should be able to see exactly how the proposal satisfies it. "
            "Preserve the section's existing strengths and stay grounded in the CONTEXT; "
            "do not invent capabilities."
        )
        return self.generate_section(spec, feedback=feedback, prior=section)

    # ------------------------------------------------------------------ #
    # Compliance matrix
    # ------------------------------------------------------------------ #
    def _compliance_evidence(self, k_per_query: int = 4) -> List[CitationEvidence]:
        """Dedup'd solicitation chunks most relevant to compliance (L/M/shall)."""
        seen: Dict[str, CitationEvidence] = {}
        for query in _COMPLIANCE_QUERIES:
            for ev in self.solicitation.search(query, k=k_per_query):
                if ev.label not in seen:
                    seen[ev.label] = ev
        return list(seen.values())

    def extract_compliance_matrix(
        self, draft: ProposalDraft, *, max_items: int = 15
    ) -> ComplianceMatrix:
        """Extract key requirements and map them to the draft's sections."""
        evidence = self._compliance_evidence()
        items: List[ComplianceItem] = []
        used_llm = False
        if evidence:
            if self.llm.available:
                items = self._extract_requirements_llm(evidence, max_items)
                used_llm = bool(items)
            if not items:
                items = self._extract_requirements_rules(evidence, max_items)
                used_llm = False
        self._map_requirements_to_sections(items, draft)
        return ComplianceMatrix(
            notice_id=draft.notice_id, items=items, used_llm=used_llm
        )

    def _extract_requirements_llm(
        self, evidence: List[CitationEvidence], max_items: int
    ) -> List[ComplianceItem]:
        labels = [e.label for e in evidence]
        context_parts = []
        for e in evidence:
            full = self.solicitation.by_label.get(e.label, e.snippet)
            context_parts.append(f"[{e.label}]\n{full[:1500]}")
        context = "\n\n".join(context_parts)
        instruction = (
            f"Extract up to {max_items} of the MOST IMPORTANT compliance "
            "requirements an offeror must satisfy, drawn ONLY from the CONTEXT. "
            "Prioritize Section L submission instructions, Section M evaluation "
            "factors, and mandatory 'shall'/'must' statements. Be concise and "
            "specific; do not invent requirements. Return a JSON object of the "
            'form {"requirements": [{"requirement": str, "category": one of '
            '["Section L","Section M","Evaluation","SOW (shall)","Other"], '
            '"source_label": one of the bracketed labels shown in the CONTEXT}]}. '
            "Use the exact source_label that the requirement came from."
        )
        try:
            result = self.llm.complete(
                instruction=instruction,
                context=context,
                sources=labels,
                max_tokens=1400,
                json_mode=True,
            )
        except LLMUnavailable as exc:
            logger.info("Compliance LLM extraction unavailable: %s", exc)
            return []

        raw = _parse_json_object(result.text)
        if not raw:
            return []
        reqs = raw.get("requirements") or raw.get("items") or []
        items: List[ComplianceItem] = []
        label_set = set(labels)
        for entry in reqs:
            if not isinstance(entry, dict):
                continue
            text = str(entry.get("requirement") or "").strip()
            if not text:
                continue
            category = _normalize_category(entry.get("category"))
            label = str(entry.get("source_label") or "").strip()
            if label not in label_set:
                # Model returned an unknown label; re-ground via semantic search.
                best = self.solicitation.search(text, k=1)
                label = best[0].label if best else (labels[0] if labels else "Solicitation")
            snippet = _snippet(self.solicitation.by_label.get(label, ""))
            items.append(
                ComplianceItem(
                    requirement=text,
                    category=category,
                    source_label=label,
                    source_snippet=snippet,
                )
            )
            if len(items) >= max_items:
                break
        return items

    def _extract_requirements_rules(
        self, evidence: List[CitationEvidence], max_items: int
    ) -> List[ComplianceItem]:
        """Regex-based fallback: pull 'shall'/'must'/evaluation sentences."""
        items: List[ComplianceItem] = []
        seen_text: set = set()
        # Evaluation/Section M language first, then general mandatory statements.
        for ev in evidence:
            full = self.solicitation.by_label.get(ev.label, ev.snippet)
            for match in _REQUIREMENT_RE.finditer(full):
                sentence = " ".join(match.group(0).split()).strip()
                if len(sentence) < 25 or len(sentence) > 400:
                    continue
                key = sentence.lower()
                if key in seen_text:
                    continue
                seen_text.add(key)
                items.append(
                    ComplianceItem(
                        requirement=sentence,
                        category=_categorize(sentence),
                        source_label=ev.label,
                        source_snippet=_snippet(full),
                    )
                )
                if len(items) >= max_items:
                    return _prioritize(items)
        return _prioritize(items)

    def remap_compliance(self, matrix: ComplianceMatrix, draft: ProposalDraft) -> None:
        """Re-derive each requirement's mapped section + status from current text.

        Preserves any status/notes the user manually edited away from the
        auto-default is *not* attempted here — call after section content changes
        (e.g. "strengthen") to refresh the auto-mapping.
        """
        self._map_requirements_to_sections(matrix.items, draft)

    def _map_requirements_to_sections(
        self, items: List[ComplianceItem], draft: ProposalDraft
    ) -> None:
        """Assign each requirement to its best-matching section + a status."""
        if not items or not draft.sections:
            return
        sections = draft.sections
        sec_texts = [f"{s.title}. {s.content}" for s in sections]
        sec_vecs = self.backend.embed(sec_texts)
        req_vecs = self.backend.embed([it.requirement for it in items])
        with np.errstate(divide="ignore", over="ignore", invalid="ignore"):
            sims = req_vecs @ sec_vecs.T
        sims = np.nan_to_num(sims, nan=-1.0, posinf=-1.0, neginf=-1.0)
        for i, item in enumerate(items):
            row = sims[i]
            j = int(np.argmax(row))
            best = float(row[j])
            item.section_id = sections[j].section_id
            item.section_title = sections[j].title
            item.match_score = best
            if best >= _ADDRESSED_THRESHOLD:
                item.status = STATUS_ADDRESSED
            elif best >= _PARTIAL_THRESHOLD:
                item.status = STATUS_PARTIAL
            else:
                item.status = STATUS_NOT_ADDRESSED


# --------------------------------------------------------------------------- #
# Post-processing & DOCX export
# --------------------------------------------------------------------------- #
_SOURCES_BLOCK_RE = re.compile(
    r"(?:\n\s*)+(?:#{1,4}\s*)?\*{0,2}Sources\*{0,2}\s*:.*",
    re.IGNORECASE | re.DOTALL,
)
_REFERENCES_BLOCK_RE = re.compile(
    r"(?:\n\s*)+(?:#{1,4}\s*)?\*{0,2}"
    r"(?:References|Citations|Bibliography|Works\s+Cited)\*{0,2}\s*:.*",
    re.IGNORECASE | re.DOTALL,
)
_MD_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_MD_ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")
_MD_UNDERSCORE_RE = re.compile(r"_(.+?)_")
_TAB_HEADING_RE = re.compile(r"^Tab\s+\d+\b", re.IGNORECASE)
_APPENDIX_HEADING_RE = re.compile(r"^Appendix\s+\d+\b", re.IGNORECASE)
_NUMBERED_SECTION_RE = re.compile(r"^\d+(?:\.\d+)+\s+\S")
_ALLCAPS_HEADING_RE = re.compile(r"^[A-Z][A-Z0-9\s/&.,'-]{6,}$")


def clean_proposal_content(text: str, *, for_export: bool = False) -> str:
    """Strip leaked Sources blocks, markdown artifacts, and export noise."""
    if not text:
        return ""
    out = text.strip()
    out = _SOURCES_BLOCK_RE.sub("", out)
    out = _REFERENCES_BLOCK_RE.sub("", out)
    out = re.sub(r"#{1,6}\s*", "", out)
    out = _MD_BOLD_RE.sub(r"\1", out)
    out = _MD_ITALIC_RE.sub(r"\1", out)
    out = _MD_UNDERSCORE_RE.sub(r"\1", out)
    if for_export:
        # Inline bracket citations belong in the Sources appendix, not body prose.
        out = re.sub(r"\[(?:KB(?:\s+style)?|Solicitation)[^\]]*\]", "", out)
    # Collapse extra whitespace left by removals.
    out = re.sub(r"[ \t]+\n", "\n", out)
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip()


def _docx_add_rich_paragraph(document, text: str, *, style: Optional[str] = None):
    """Add a paragraph, converting any remaining **bold** markers to Word runs."""
    para = document.add_paragraph(style=style)
    if not text:
        return para
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)
    return para


def _docx_add_title_page(document, text: str) -> None:
    """Render a formal title page matching the uploaded proposal template."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    emphasis = {
        "title page", "volume 1", "technical proposal",
    }
    for i, line in enumerate(lines):
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        low = line.lower()
        if i == 0 or low in emphasis:
            run.bold = True
        if i == 0:
            run.font.size = Pt(16)
        elif low == "technical proposal":
            run.font.size = Pt(14)
    document.add_page_break()


def _docx_add_table_of_contents(document, text: str) -> None:
    """Render the Table of Contents in the traditional Tab / Appendix list style."""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return
    document.add_heading(lines[0], level=1)
    for line in lines[1:]:
        if _TAB_HEADING_RE.match(line) or _APPENDIX_HEADING_RE.match(line):
            para = document.add_paragraph()
            run = para.add_run(line)
            run.bold = True
        else:
            document.add_paragraph(line)
    document.add_page_break()


def _docx_add_proposal_body(document, text: str) -> None:
    """Add cleaned proposal prose with heading detection for Tabs and subsections."""
    blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
    for block in blocks:
        lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
        if not lines:
            continue
        if len(lines) == 1:
            line = lines[0]
            if _TAB_HEADING_RE.match(line):
                document.add_heading(line, level=1)
            elif _APPENDIX_HEADING_RE.match(line) or _NUMBERED_SECTION_RE.match(line):
                document.add_heading(line, level=2)
            elif _ALLCAPS_HEADING_RE.match(line) and len(line) < 80:
                document.add_heading(line, level=2)
            elif line.startswith(("- ", "• ", "* ")):
                document.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                _docx_add_rich_paragraph(document, line)
            continue
        # Multi-line block: first line may be a heading.
        first = lines[0]
        rest = lines[1:]
        if _TAB_HEADING_RE.match(first):
            document.add_heading(first, level=1)
        elif (
            _APPENDIX_HEADING_RE.match(first)
            or _NUMBERED_SECTION_RE.match(first)
            or (_ALLCAPS_HEADING_RE.match(first) and len(first) < 80)
        ):
            document.add_heading(first, level=2)
        else:
            _docx_add_rich_paragraph(document, first)
        for line in rest:
            if line.startswith(("- ", "• ", "* ")):
                document.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                _docx_add_rich_paragraph(document, line)


def _docx_add_sources_appendix(
    document,
    sections: List[ProposalSection],
) -> None:
    """Append a dedicated sources section — not inline in proposal Tabs."""
    sourced = [(s.title, s.sources) for s in sections if s.sources]
    if not sourced:
        return
    document.add_page_break()
    document.add_heading("Document Sources and References", level=1)
    document.add_paragraph(
        "The following sources grounded each section of this draft. "
        "Review and verify before submission.",
        style="Intense Quote",
    )
    for title, sources in sourced:
        document.add_heading(title, level=2)
        for src in sources:
            label = src.label
            if src.origin == "style_exemplar":
                label = f"{label} (writing-style template)"
            document.add_paragraph(label, style="List Bullet")


def draft_to_docx_bytes(draft: ProposalDraft) -> bytes:
    """Render a :class:`ProposalDraft` to a professional .docx proposal volume."""
    import docx

    document = docx.Document()
    body_sections: List[ProposalSection] = []

    for section in draft.sections:
        cleaned = clean_proposal_content(section.content, for_export=True)
        if section.section_id == "title_page":
            _docx_add_title_page(document, cleaned)
            continue
        if section.section_id == "table_of_contents":
            _docx_add_table_of_contents(document, cleaned)
            continue
        if section.section_id.startswith("tab_"):
            document.add_page_break()
        # Skip duplicate heading when the Tab title is already the first line.
        starts_with_tab = bool(cleaned) and _TAB_HEADING_RE.match(cleaned.splitlines()[0])
        if not starts_with_tab:
            document.add_heading(section.title, level=1)
        _docx_add_proposal_body(document, cleaned)
        body_sections.append(section)

    _docx_add_sources_appendix(document, body_sections)

    if draft.compliance and draft.compliance.items:
        _add_compliance_table_to_doc(document, draft.compliance)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_compliance_table_to_doc(document, matrix: ComplianceMatrix) -> None:
    """Append the compliance matrix as a Word table."""
    document.add_page_break()
    document.add_heading("Compliance Matrix", level=1)
    counts = matrix.status_counts()
    document.add_paragraph(
        f"Summary: {counts[STATUS_ADDRESSED]} addressed · "
        f"{counts[STATUS_PARTIAL]} partial · "
        f"{counts[STATUS_NOT_ADDRESSED]} not addressed. "
        "Status is an auto-generated starting point — review and adjust.",
        style="Intense Quote",
    )
    headers = ["#", "Requirement / Factor", "Category", "Source", "Section", "Status", "Notes"]
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:  # pragma: no cover - style availability varies
        pass
    for cell, head in zip(table.rows[0].cells, headers):
        cell.text = head
    for i, item in enumerate(matrix.items, start=1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = item.requirement
        cells[2].text = item.category
        cells[3].text = item.source_label
        cells[4].text = item.section_title or "—"
        cells[5].text = item.status
        cells[6].text = item.notes or ""


def compliance_matrix_to_xlsx_bytes(matrix: ComplianceMatrix) -> bytes:
    """Render the compliance matrix to an .xlsx workbook (bytes)."""
    import pandas as pd

    rows = [
        {
            "#": i,
            "Requirement / Factor": item.requirement,
            "Category": item.category,
            "Source": item.source_label,
            "Source Excerpt": item.source_snippet,
            "Proposal Section": item.section_title or "",
            "Status": item.status,
            "Match Score": round(item.match_score, 3) if item.match_score is not None else "",
            "Notes": item.notes or "",
        }
        for i, item in enumerate(matrix.items, start=1)
    ]
    df = pd.DataFrame(
        rows,
        columns=[
            "#", "Requirement / Factor", "Category", "Source", "Source Excerpt",
            "Proposal Section", "Status", "Match Score", "Notes",
        ],
    )
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Compliance Matrix")
        worksheet = writer.sheets["Compliance Matrix"]
        widths = {"A": 5, "B": 60, "C": 14, "D": 32, "E": 50, "F": 26, "G": 14, "H": 12, "I": 40}
        for col, width in widths.items():
            worksheet.column_dimensions[col].width = width
    return buffer.getvalue()
